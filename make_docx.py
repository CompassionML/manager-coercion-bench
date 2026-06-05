"""
Convert REPORT.md -> report.docx with the two figures embedded inline.

Google Docs imports .docx natively (File > Open, or just drag into Drive),
converting it to an editable Google Doc with the images preserved — which a
raw Markdown file cannot do.

Lightweight Markdown handling tuned to REPORT.md (each block is on its own
line): headings (#..####), bullet/numbered lists, blockquotes, inline
**bold**/*italic*/`code`, and ![alt](path) images.
"""

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SRC = "REPORT.md"
OUT = "report.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
IMG = re.compile(r"^!\[(.*)\]\((.*)\)\s*$")
HEAD = re.compile(r"^(#{1,6})\s+(.*)$")
NUM = re.compile(r"^(\d+)\.\s+(.*)$")


def add_inline(p, text):
    text = text.replace(r"\*", "*")
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"
        else:
            p.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = open(SRC, encoding="utf-8").read().splitlines()
    for line in lines:
        s = line.rstrip()
        if not s.strip():
            continue
        if s.strip() == "---":
            continue

        mimg = IMG.match(s)
        if mimg:
            alt, path = mimg.group(1), mimg.group(2)
            width = Inches(5.6) if "tactics" in path else Inches(6.2)
            doc.add_picture(path, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(alt)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue

        mhead = HEAD.match(s)
        if mhead:
            level = len(mhead.group(1))
            doc.add_heading(mhead.group(2).strip(), level=0 if level == 1 else level - 1)
            continue

        if s.startswith("> "):
            add_inline(doc.add_paragraph(style="Quote"), s[2:])
            continue
        if s.startswith("- "):
            add_inline(doc.add_paragraph(style="List Bullet"), s[2:])
            continue
        mnum = NUM.match(s)
        if mnum:
            add_inline(doc.add_paragraph(style="List Number"), mnum.group(2))
            continue

        add_inline(doc.add_paragraph(), s)

    doc.save(OUT)

    # verify
    chk = Document(OUT)
    n_img = len(chk.inline_shapes)
    print(f"wrote {OUT}  |  paragraphs={len(chk.paragraphs)}  inline_images={n_img}")


if __name__ == "__main__":
    main()
